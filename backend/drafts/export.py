"""
Renders a Draft's title/content into a downloadable file - PDF via PyMuPDF,
DOCX via python-docx. Both libraries are already dependencies of this
project (used for reading uploaded PDFs/DOCX in rag/document_processor.py),
so this doesn't add anything new to requirements.
"""
import io

import fitz
from docx import Document
from docx.shared import Pt


def build_docx_bytes(title: str, content: str) -> bytes:
    document = Document()
    document.add_heading(title or "Untitled Draft", level=1)

    for line in content.split("\n"):
        paragraph = document.add_paragraph(line)
        for run in paragraph.runs:
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _wrap_paragraph(paragraph: str, max_width: float, fontsize: int, fontname: str) -> list:
    """
    Word-wraps one paragraph into lines that fit max_width, measuring each
    candidate line with the real font metrics (fitz.get_text_length) rather
    than guessing by character count - avoids the all-or-nothing failure
    mode of insert_textbox (which draws nothing at all if a block of text
    doesn't fit its box) by placing each already-measured line individually
    with insert_text instead.
    """
    words = paragraph.split(" ")
    lines = []
    current = ""

    for word in words:
        candidate = f"{current} {word}".strip()
        if fitz.get_text_length(candidate, fontname=fontname, fontsize=fontsize) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word

    if current:
        lines.append(current)

    return lines or [""]


def build_pdf_bytes(title: str, content: str) -> bytes:
    margin = 50
    font_size = 11
    title_font_size = 16
    fontname = "helv"
    line_height = font_size * 1.4

    doc = fitz.open()
    page = doc.new_page()
    max_width = page.rect.width - 2 * margin
    bottom_limit = page.rect.height - margin

    y = margin + title_font_size
    page.insert_text((margin, y), title or "Untitled Draft", fontsize=title_font_size, fontname=fontname)
    y += line_height * 2

    for paragraph in content.split("\n"):
        if not paragraph.strip():
            y += line_height
            continue

        for line in _wrap_paragraph(paragraph, max_width, font_size, fontname):
            if y > bottom_limit:
                page = doc.new_page()
                y = margin

            page.insert_text((margin, y), line, fontsize=font_size, fontname=fontname)
            y += line_height

        y += line_height * 0.5

    buffer = io.BytesIO()
    doc.save(buffer)
    doc.close()
    return buffer.getvalue()
